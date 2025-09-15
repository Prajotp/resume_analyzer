from flask import Flask, render_template, request, flash, redirect, url_for
import PyPDF2
import docx
import re
import io
import os
from datetime import datetime
from collections import defaultdict

app = Flask(__name__)
app.secret_key = 'resume-analyzer-secret-key-2023'
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB max file size

# Job roles database
JOB_ROLES = {
    "Software Developer": {
        "department": "IT & Software",
        "skills": ["python", "java", "javascript", "c++", "sql", "git", "agile", "rest api", "docker", "kubernetes", "angular", "html", "css", "bootstrap"],
        "keywords": ["software", "developer", "programming", "coding", "backend", "frontend", "full stack", "web application"],
        "min_experience": 0,
        "senior_level": 5
    },
    "Frontend Developer": {
        "department": "IT & Software",
        "skills": ["html", "css", "javascript", "angular", "bootstrap", "sass", "responsive design", "ui/ux"],
        "keywords": ["frontend", "web development", "user interface", "single page application", "spa"],
        "min_experience": 1,
        "senior_level": 5
    },
    "Web Developer": {
        "department": "IT & Software",
        "skills": ["html", "css", "javascript", "angular", "bootstrap", "java", "mysql", "rest api"],
        "keywords": ["web developer", "web application", "website", "web development"],
        "min_experience": 1,
        "senior_level": 5
    },
    "Angular Developer": {
        "department": "IT & Software",
        "skills": ["angular", "typescript", "javascript", "html", "css", "bootstrap", "rxjs", "ngrx"],
        "keywords": ["angular", "frontend", "single page application", "spa", "web development"],
        "min_experience": 1,
        "senior_level": 5
    },
    "Java Developer": {
        "department": "IT & Software",
        "skills": ["java", "spring", "hibernate", "j2ee", "rest api", "sql", "microservices"],
        "keywords": ["java", "backend", "enterprise", "spring framework"],
        "min_experience": 1,
        "senior_level": 5
    }
}

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        file.seek(0)  # Reset file pointer
        doc = docx.Document(io.BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def extract_text_from_file(file):
    """Extract text from uploaded file based on extension"""
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file)
    else:
        return None

def calculate_experience(text):
    """Calculate total years of experience from resume text"""
    text_lower = text.lower()
    
    # First, look for explicitly stated experience
    explicit_experience_patterns = [
        r'(\d+\.?\d*)\s*\+?\s*years?[\s\w]*experience',
        r'experience.*?(\d+\.?\d*)\s*\+?\s*years?',
        r'(\d+\.?\d*)\s*\+?\s*years?.*?professional',
        r'(\d+\.?\d*)\s*\+?\s*years?.*?work',
        r'(\d+\.?\d*)\s*\+?\s*years?.*?industry'
    ]
    
    found_experience = 0
    for pattern in explicit_experience_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            if isinstance(match, tuple):
                match = match[0]
            try:
                years = float(match)
                if years > found_experience:
                    found_experience = years
                    return round(found_experience, 1)  # Return immediately if found
            except ValueError:
                pass
    
    # If no explicit experience found, calculate from employment dates
    date_pattern = r'(\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|[0-9]{1,2}/)?\s*\d{4})\s*[-–—]\s*(\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|[0-9]{1,2}/)?\s*\d{4}|present|current|now)'
    date_ranges = re.findall(date_pattern, text_lower, re.IGNORECASE)
    
    total_months = 0
    current_year = datetime.now().year
    current_month = datetime.now().month
    
    for start_str, end_str in date_ranges:
        try:
            # Extract year from start date
            start_year_match = re.search(r'(\d{4})', start_str)
            if start_year_match:
                start_year = int(start_year_match.group(1))
                
                # Try to extract month
                start_month = 1  # Default to January
                month_match = re.search(r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)', start_str, re.IGNORECASE)
                if month_match:
                    month_str = month_match.group(1).lower()
                    month_map = {
                        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
                    }
                    start_month = month_map.get(month_str, 1)
            else:
                continue
            
            # Extract year from end date
            if end_str.lower() in ['present', 'current', 'now']:
                end_year = current_year
                end_month = current_month
            else:
                end_year_match = re.search(r'(\d{4})', end_str)
                if end_year_match:
                    end_year = int(end_year_match.group(1))
                    end_month = 12  # Default to December
                    
                    # Try to extract month
                    month_match = re.search(r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)', end_str, re.IGNORECASE)
                    if month_match:
                        month_str = month_match.group(1).lower()
                        month_map = {
                            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
                        }
                        end_month = month_map.get(month_str, 12)
                else:
                    continue
            
            # Calculate months of experience for this range
            months = (end_year - start_year) * 12 + (end_month - start_month) + 1
            total_months += months
        except:
            continue
    
    # Convert months to years
    if total_months > 0:
        found_experience = total_months / 12
    
    return round(found_experience, 1)

def extract_education(text):
    """Extract education information from resume text"""
    text_lower = text.lower()
    
    education_info = {
        "highest_degree": "Not specified",
        "degrees": [],
        "institutions": [],
        "years": []
    }
    
    # Education section pattern
    education_section_pattern = r'(education|academic background|qualifications|academic qualifications)(.*?)(?=experience|skills|projects|$|certifications)'
    education_section_match = re.search(education_section_pattern, text_lower, re.IGNORECASE | re.DOTALL)
    
    if education_section_match:
        education_text = education_section_match.group(2)
    else:
        education_text = text_lower
    
    # Degree patterns
    degree_patterns = {
        "PhD": r'\b(ph\.?d|doctorate|doctoral)\b',
        "Master's": r'\b(master|m\.?s|m\.?sc|mba|m\.?a|m\.?ed|post graduate|pg|post graduation)\b',
        "Bachelor's": r'\b(bachelor|b\.?s|b\.?sc|b\.?a|b\.?com|btech|b\.?tech|be|b\.?e|engineering)\b',
        "Associate": r'\b(associate|a\.?a|a\.?sc|a\.?as)\b',
        "Diploma": r'\b(diploma|certificate|certification)\b'
    }
    
    # Find highest degree
    highest_degree = "Not specified"
    for degree, pattern in degree_patterns.items():
        if re.search(pattern, education_text, re.IGNORECASE):
            highest_degree = degree
            break
    
    # Extract degree details
    degree_details = []
    degree_line_pattern = r'([A-Za-z\s]+?)\s*\|\s*([A-Za-z\s]+?)\s*\|\s*(\d{4})'
    matches = re.findall(degree_line_pattern, text, re.IGNORECASE)
    
    for match in matches:
        degree, institution, year = match
        degree_details.append({
            "degree": degree.strip(),
            "institution": institution.strip(),
            "year": year.strip()
        })
    
    # If no structured format found, try to extract from text
    if not degree_details:
        # Look for education lines
        education_lines = re.findall(r'([A-Za-z\s]+?)\s*(?:in|of|,)\s*([A-Za-z\s]+?)\s*(?:\d{4})', education_text, re.IGNORECASE)
        for line in education_lines:
            if len(line) >= 2:
                degree_details.append({
                    "degree": line[0].strip(),
                    "institution": line[1].strip(),
                    "year": "Not specified"
                })
    
    education_info["highest_degree"] = highest_degree
    education_info["degrees"] = degree_details
    
    return education_info

def analyze_resume(text):
    """Analyze resume text and extract key information"""
    text_lower = text.lower()
    
    # Extract skills with better matching
    found_skills = []
    all_skills = []
    for job in JOB_ROLES.values():
        all_skills.extend(job["skills"])
    
    # Remove duplicates and sort
    all_skills = sorted(list(set(all_skills)))
    
    for skill in all_skills:
        # Use word boundary matching for better accuracy
        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
            if skill not in found_skills:
                found_skills.append(skill)
    
    # Calculate experience
    experience_years = calculate_experience(text)
    
    # Extract education information
    education_info = extract_education(text)
    
    # Extract contact information
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    email = email_match.group(0) if email_match else "Not found"
    
    phone_match = re.search(r'(\+?(\d{1,3})?[\s-]?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4})', text)
    phone = phone_match.group(0) if phone_match else "Not found"
    
    # Extract certifications with better patterns
    certifications = []
    cert_patterns = [
        r'\b([A-Za-z\s]+)\s+(?:certification|certified)\b',
        r'\b(?:certification|certificate)\s+in\s+([A-Za-z\s]+)\b',
        r'\b([A-Za-z]+)\s+(?:cert|certification)\b'
    ]
    
    for pattern in cert_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            if isinstance(match, tuple):
                match = match[0]
            if len(match.strip()) > 3:  # Filter out short meaningless matches
                certifications.append(match.strip())
    
    # Extract job titles/positions
    positions = []
    position_patterns = [
        r'\b(?:senior|junior|lead|principal)?\s*([a-z\s]+)\s*(?:engineer|developer|analyst|manager|designer|specialist)\b',
        r'\b(?:worked as|position|role of|as a)\s+([a-z\s]+)\b'
    ]
    
    for pattern in position_patterns:
        matches = re.findall(pattern, text_lower)
        positions.extend([m.strip() for m in matches if len(m.strip()) > 3])
    
    # Calculate job suitability scores with improved algorithm
    job_scores = {}
    for job, requirements in JOB_ROLES.items():
        # Skill matching with partial matches
        skill_score = 0
        for skill in requirements["skills"]:
            if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
                skill_score += 2  # Full match
            elif skill.lower() in text_lower:
                skill_score += 1  # Partial match
        
        # Keyword matching
        keyword_score = 0
        for keyword in requirements["keywords"]:
            if re.search(r'\b' + re.escape(keyword.lower()) + r'\b', text_lower):
                keyword_score += 1
        
        # Experience factor (more nuanced)
        exp_factor = 1.0
        if experience_years >= requirements["min_experience"]:
            if experience_years >= requirements["senior_level"]:
                exp_factor = 1.3  # Bonus for senior-level experience
            else:
                exp_factor = 1.1  # Standard bonus for meeting minimum
        else:
            exp_factor = 0.8  # Small penalty for insufficient experience
        
        # Calculate final score
        total_score = (skill_score + keyword_score) * exp_factor
        max_possible = (len(requirements["skills"]) * 2 + len(requirements["keywords"])) * 1.3
        percentage = int((total_score / max_possible) * 100) if max_possible > 0 else 0
        
        job_scores[job] = min(percentage, 100)
    
    # Determine eligibility (threshold: 25% match)
    eligible_jobs = {job: score for job, score in job_scores.items() if score >= 25}
    best_match = max(job_scores.items(), key=lambda x: x[1]) if job_scores else ("None", 0)
    
    # Group jobs by department
    jobs_by_department = defaultdict(dict)
    for job, score in eligible_jobs.items():
        dept = JOB_ROLES[job]["department"]
        jobs_by_department[dept][job] = score
    
    # Determine experience level
    if experience_years >= 10:
        experience_level = "Senior/Executive"
    elif experience_years >= 5:
        experience_level = "Mid-Senior"
    elif experience_years >= 2:
        experience_level = "Mid-Level"
    elif experience_years > 0:
        experience_level = "Entry-Level"
    else:
        experience_level = "No Experience"
    
    return {
        "skills": found_skills,
        "experience": experience_years,
        "experience_level": experience_level,
        "education": education_info,
        "certifications": certifications[:10],  # Limit to top 10
        "positions": positions[:5],  # Limit to top 5 positions
        "contact": {
            "email": email,
            "phone": phone
        },
        "job_scores": job_scores,
        "eligible_jobs": eligible_jobs,
        "jobs_by_department": dict(jobs_by_department),
        "best_match": best_match,
        "raw_text": text  # For debugging
    }

@app.route('/')
def index():
    departments = sorted({job["department"] for job in JOB_ROLES.values()})
    return render_template('index.html', departments=departments)

@app.route('/analyze', methods=['POST'])
def analyze():
    if 'resume' not in request.files:
        flash('Please upload a resume file', 'error')
        return redirect(url_for('index'))
    
    file = request.files['resume']
    
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('index'))
    
    # Check file type
    if not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        flash('Please upload a PDF or DOCX file', 'error')
        return redirect(url_for('index'))
    
    # Extract text from file
    text = extract_text_from_file(file)
    if text is None:
        flash('Error reading file. Please try another file.', 'error')
        return redirect(url_for('index'))
    
    if len(text.strip()) < 50:
        flash('The file appears to have very little text. Please upload a valid resume.', 'error')
        return redirect(url_for('index'))
    
    # Analyze the resume
    analysis = analyze_resume(text)
    
    return render_template('results.html', analysis=analysis, jobs=JOB_ROLES)

@app.errorhandler(413)
def too_large(e):
    flash('File is too large. Maximum size is 5MB.', 'error')
    return redirect(url_for('index'))

@app.errorhandler(500)
def internal_error(error):
    flash('An internal error occurred. Please try again.', 'error')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)